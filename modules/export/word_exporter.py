import os
import logging
from .base_exporter import BaseExporter

logger = logging.getLogger(__name__)

class WordExporter(BaseExporter):
    """Word文档导出器"""
    
    def export_questions(self, questions, output_path=None, include_answers=False):
        """
        导出题目到Word文档
        
        Args:
            questions: 题目数据列表
            output_path: 输出文件路径，如果为None则使用默认路径
            include_answers: 是否包含答案
            
        Returns:
            输出文件路径
        """
        try:
            # 如果未指定输出路径，使用默认路径
            if output_path is None:
                file_suffix = "题目和答案" if include_answers else "题目"
                output_path = self.get_default_filename(f"错题_{file_suffix}", "docx")
            
            # 使用python-docx导出Word文档
            from docx import Document
            from docx.shared import RGBColor
            
            doc = Document()
            
            # 添加标题
            title_text = "错题集" if not include_answers else "错题集（含答案）"
            title = doc.add_heading(title_text, 0)
            title.alignment = 1  # 居中对齐
            
            # 遍历每道题目
            for i, question in enumerate(questions, 1):
                # 题目标题
                heading_text = f"{i}. [{question['subject']}] [{question['question_type']}] [难度: {'★' * question['difficulty']}]"
                heading = doc.add_heading(heading_text, level=2)
                
                # 设置标题颜色为蓝色
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 102, 204)
                
                # 题目内容
                doc.add_paragraph(question["question_text"])
                
                # 如果包含答案
                if include_answers:
                    # 答案标题
                    answer_heading = doc.add_heading("正确答案:", level=3)
                    for run in answer_heading.runs:
                        run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                    
                    # 答案内容
                    doc.add_paragraph(question.get("correct_answer", ""))
                    
                    # 如果有解析
                    if question.get("explanation"):
                        # 解析标题
                        explanation_heading = doc.add_heading("解析:", level=3)
                        for run in explanation_heading.runs:
                            run.font.color.rgb = RGBColor(70, 130, 180)  # 钢青色
                        
                        # 解析内容
                        doc.add_paragraph(question.get("explanation", ""))
                
                # 添加分隔线（除了最后一题）
                if i < len(questions):
                    doc.add_paragraph("_" * 50)
            
            # 保存文档
            doc.save(output_path)
            
            return output_path
            
        except ImportError:
            logger.error("导出Word需要安装python-docx库")
            raise ImportError("导出Word需要安装python-docx库，请使用命令：pip install python-docx")
        except Exception as e:
            logger.error(f"导出Word失败: {str(e)}")
            raise