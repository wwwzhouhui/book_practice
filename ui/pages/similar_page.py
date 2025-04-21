import gradio as gr
import pandas as pd
from modules.storage.database import Database
from modules.ai_generator.question_generator import QuestionGenerator
import logging
from datetime import datetime
import os
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import configparser
import platform
import traceback

logger = logging.getLogger(__name__)

def create_similar_page():
    """创建同类题生成页面"""
    
    # 初始化数据库和生成器
    db = Database()
    generator = QuestionGenerator()
    
    # 确保导出目录存在
    export_dir = os.path.join(os.getcwd(), "exports")
    os.makedirs(export_dir, exist_ok=True)
    
    with gr.Tab("生成同类题"):
        with gr.Row():
            with gr.Column(scale=1):
                # 筛选条件
                subject = gr.Dropdown(
                    choices=["数学", "物理", "化学"],
                    label="学科",
                    value="数学"
                )
                question_type = gr.Dropdown(
                    choices=["","选择题", "填空题", "判断题", "问答题"],
                    label="题型"
                )
                # difficulty = gr.Dropdown(
                #     choices=["", "简单", "中等", "困难"],
                #     label="难度",
                #     value=""
                # )
                date_range = gr.Dropdown(
                    choices=["全部", "最近一周", "最近一月"],
                    label="时间范围",
                    value="全部"
                )
                search_btn = gr.Button("查询错题")
            
            with gr.Column(scale=2):
                # 根据官方文档修改表格配置
                question_table = gr.Dataframe(
                    headers=["题目ID", "题目内容", "学科", "题型", "难度", "创建时间"],
                    datatype=["number", "str", "str", "str", "str", "str"],
                    interactive=False,  # 设置为False，因为我们只需要选择功能
                    label="错题列表",
                    value=None,  # 初始值为空
                    wrap=True,
                    row_count=10
                )
                
                with gr.Row():
                    selected_count = gr.Text(label="已选题目ID", value="", interactive=False)
                    gen_count = gr.Slider(
                        minimum=1,
                        maximum=5,
                        value=1,
                        step=1,
                        label="生成数量"
                    )
                    generate_btn = gr.Button("生成同类题",interactive=False)

                # 生成结果
                with gr.Row():
                    # 添加生成结果表格
                    generated_table = gr.Dataframe(
                        headers=["题目内容", "答案", "解析"],
                        label="生成结果",
                        wrap=True,
                        row_count=5,
                        interactive=False
                    )
                    
                with gr.Row():
                    result_status = gr.Text(label="生成状态", interactive=False)
                    save_btn = gr.Button("保存生成结果", interactive=False)
                    export_questions_btn = gr.Button("导出题目", interactive=False)  # 修改按钮
                    export_answers_btn = gr.Button("导出答案", interactive=False)   # 新增按钮
                    
                # 添加下载文件组件
                download_file = gr.File(label="下载文件", visible=False, interactive=True)

        # 定义事件处理函数
        def search_questions(subject, question_type, date_range):
            """查询错题"""
            try:
                # 根据 Database 类的实际方法参数进行查询
                questions = db.get_all_error_questions(
                    subject=subject,
                    question_type=question_type if question_type else None,
                    time_range=date_range
                )
                
                if not questions:
                    return gr.update(value=pd.DataFrame()), "未找到符合条件的错题"
                
                # 在内存中根据难度进行过滤
                # if difficulty:
                #     questions = [q for q in questions if q["difficulty"] == difficulty]
                
                df = pd.DataFrame({
                    "题目ID": [q["id"] for q in questions],
                    "题目内容": [q["question_text"] for q in questions],
                    "学科": [q["subject"] for q in questions],
                    "题型": [q["question_type"] for q in questions],
                    "难度": [q["difficulty"] for q in questions],
                    "创建时间": [q["created_at"] for q in questions]
                })
                
                return gr.update(value=df, interactive=True), ""
                
            except Exception as e:
                logger.error(f"查询错题失败: {str(e)}")
                return gr.update(value=pd.DataFrame()), f"查询失败: {str(e)}"

        def on_select(evt: gr.SelectData, df):
            """处理表格选择事件"""
            try:
                # 获取选中的行索引
                if isinstance(evt.index, tuple) or isinstance(evt.index, list):
                    # 如果是元组或列表格式 [row, col]，只取行号
                    selected_indices = evt.index[0]
                else:
                    # 如果是单个整数
                    selected_indices = evt.index
                
                # 如果df为空，直接返回
                if df is None or len(df) == 0:
                    return "", False
                    
                # 获取选中行的题目ID
                selected_id = str(df.iloc[selected_indices]["题目ID"])
                
                logger.info(f"选中了行，题目ID: {selected_id}")
                
                # 返回选中的ID字符串和生成按钮状态
                return selected_id, gr.update(interactive=True)
                
            except Exception as e:
                logger.error(f"处理选择事件失败: {str(e)}")
                return "", False

        def generate_similar_questions(selected_count, gen_count):
            """生成同类题"""
            try:
                # 获取选中的错题
                question_ids = selected_count
                source_questions = [db.get_error_question(qid) for qid in question_ids]
                
                # 生成同类题
                generated = generator.generate_similar_questions(
                    source_questions,
                    count=gen_count
                )
                
                if not generated:
                    return (
                        "生成失败", 
                        gr.update(value=None), 
                        gr.update(interactive=False),  # save_btn
                        gr.update(interactive=False),  # export_questions_btn
                        gr.update(interactive=False)   # export_answers_btn
                    )
                
                # 转换为DataFrame
                df = pd.DataFrame({
                    "题目内容": [q["question_text"] for q in generated],
                    "答案": [q["answer"] for q in generated],
                    "解析": [q["explanation"] for q in generated]
                })
                
                return (
                    "生成成功", 
                    gr.update(value=df), 
                    gr.update(interactive=True),  # save_btn
                    gr.update(interactive=True),  # export_questions_btn
                    gr.update(interactive=True)   # export_answers_btn
                )
                
            except Exception as e:
                logger.error(f"生成同类题失败: {str(e)}")
                return (
                    f"生成失败: {str(e)}", 
                    gr.update(value=None), 
                    gr.update(interactive=False),  # save_btn
                    gr.update(interactive=False),  # export_questions_btn
                    gr.update(interactive=False)   # export_answers_btn
                )

        def save_generated_questions(result_df, selected_count):
            """保存生成的题目"""
            try:
                if result_df is None or len(result_df) == 0:
                    return "没有可保存的题目"
                
                # 获取源题目信息
                source_id = selected_count[0] if selected_count else None
                if not source_id:
                    return "未找到源题目ID"
                
                source_question = db.get_error_question(source_id)
                if not source_question:
                    return "未找到源题目信息"
                
                success_count = 0
                # 保存到数据库
                for _, row in result_df.iterrows():
                    question_data = {
                        "question_text": row["题目内容"],
                        "answer": row["答案"],
                        "explanation": row["解析"],
                        "subject": source_question["subject"],  # 从源题目获取学科
                        "question_type": source_question["question_type"],  # 从源题目获取题型
                        "difficulty": source_question["difficulty"],  # 从源题目获取难度
                        "source_question_id": source_id  # 设置源题目ID
                    }
                    
                    try:
                        db.save_similar_question(question_data)
                        success_count += 1
                    except Exception as e:
                        logger.error(f"保存单个题目失败: {str(e)}")
                
                return f"成功保存 {success_count} 道题目"
                
            except Exception as e:
                logger.error(f"保存生成题目失败: {str(e)}")
                return f"保存失败: {str(e)}"

        def export_to_word(result_df, export_type="questions"):
            """导出生成结果到Word文档"""
            try:
                if result_df is None or len(result_df) == 0:
                    return "没有可导出的数据", None, gr.update(visible=False)

                # 创建时间戳
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                file_prefix = "题目" if export_type == "questions" else "答案"
                
                try:
                    # 使用python-docx导出Word文档
                    from docx import Document
                    from docx.shared import RGBColor

                    doc_path = os.path.join(export_dir, f"{file_prefix}_{timestamp}.docx")
                    doc = Document()

                    # 添加标题
                    title = doc.add_heading(f"相似题目集 - {file_prefix}", 0)
                    title.alignment = 1  # 居中对齐

                    # 遍历每道题目
                    for i, row in result_df.iterrows():
                        if export_type == "questions":
                            # 导出题目部分
                            # 题目标题
                            heading = doc.add_heading(f"题目 {i+1}", level=2)

                            # 设置标题颜色为蓝色
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 102, 204)

                            # 题目内容
                            doc.add_paragraph(row["题目内容"])
                        else:
                            # 导出答案部分
                            # 标题
                            heading = doc.add_heading(f"题目 {i+1} 答案:", level=2)

                            # 设置标题颜色为蓝色
                            for run in heading.runs:
                                run.font.color.rgb = RGBColor(0, 102, 204)

                            # 题目内容
                            doc.add_paragraph("题目:")
                            doc.add_paragraph(row["题目内容"])

                            # 答案
                            answer_para = doc.add_paragraph()
                            answer_run = answer_para.add_run("答案: ")
                            answer_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                            answer_para.add_run(row["答案"])

                            # 解析
                            if pd.notna(row["解析"]):
                                explanation_para = doc.add_paragraph()
                                explanation_run = explanation_para.add_run("解析: ")
                                explanation_run.font.color.rgb = RGBColor(70, 130, 180)  # 钢青色
                                explanation_para.add_run(row["解析"])

                        # 添加分隔线（除了最后一题）
                        if i < len(result_df) - 1:
                            doc.add_paragraph("_" * 50)

                    # 保存文档
                    doc.save(doc_path)
                    success_msg = f"✅ {file_prefix}Word文档生成成功，请点击下载"
                    return success_msg, doc_path, gr.update(visible=True)

                except ImportError:
                    return "❌ 导出Word需要安装python-docx库，请使用命令：pip install python-docx", None, gr.update(visible=False)
                except Exception as e:
                    logger.error(f"Word导出失败: {str(e)}")
                    logger.error(traceback.format_exc())
                    return f"❌ Word导出失败: {str(e)}", None, gr.update(visible=False)

            except Exception as e:
                logger.error(f"导出失败: {str(e)}")
                logger.error(traceback.format_exc())
                return f"❌ 导出失败: {str(e)}", None, gr.update(visible=False)

        # 绑定事件处理函数
        search_btn.click(
            search_questions,
            inputs=[subject, question_type, date_range],
            outputs=[question_table, result_status]
        )
        
        question_table.select(
            fn=on_select,
            inputs=[question_table],
            outputs=[selected_count, generate_btn]  # 保持不变，但现在会根据 selected_count 是否有值来激活按钮
        )

        generate_btn.click(
            generate_similar_questions,
            inputs=[selected_count, gen_count],
            outputs=[result_status, generated_table, save_btn, export_questions_btn, export_answers_btn]
        )
        
        save_btn.click(
            save_generated_questions,
            inputs=[generated_table, selected_count],  # 添加 selected_count 作为输入
            outputs=[result_status]
        )
        
        # 导出题目按钮事件绑定
        export_questions_btn.click(
            lambda df: export_to_word(df, "questions"),
            inputs=[generated_table],
            outputs=[result_status, download_file, download_file]
        )
        
        # 导出答案按钮事件绑定
        export_answers_btn.click(
            lambda df: export_to_word(df, "answers"),
            inputs=[generated_table],
            outputs=[result_status, download_file, download_file]
        )

    return "生成同类题"








































