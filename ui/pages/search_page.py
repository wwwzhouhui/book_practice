
import gradio as gr
import pandas as pd
import numpy as np
from modules.storage.database import Database
import os
import csv
import time
import logging
import traceback

# 配置日志
logger = logging.getLogger(__name__)

def create_search_page():
    """创建错题查询页面"""
    # 初始化数据库
    db = Database()

    with gr.Column() as search_page:
        gr.Markdown("## 错题查询")

        # 搜索条件区域
        with gr.Row():
            with gr.Column(scale=2):
                search_input = gr.Textbox(
                    label="搜索关键词",
                    placeholder="输入题号或题目内容关键词..."
                )

            with gr.Column(scale=3):
                with gr.Row():
                    filters = gr.CheckboxGroup(
                        choices=["数学", "物理", "化学"],
                        label="科目筛选",
                        value=[]
                    )
                    date_range = gr.Dropdown(
                        choices=["今天", "最近一周", "最近一月", "全部"],
                        label="时间范围",
                        value="全部"
                    )

                search_button = gr.Button("查询错题", variant="primary")

        # 添加查询状态指示器
        query_status = gr.Markdown('准备就绪，点击"显示所有错题"按钮加载数据')

        # 错题列表区域
        gr.Markdown("### 错题列表")

        # 错题列表表格
        results = gr.DataFrame(
            headers=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"],
            label="查询结果",
            wrap=True
        )

        # 操作按钮区
        with gr.Row():
            with gr.Column(scale=3):
                refresh_button = gr.Button("显示所有错题", variant="primary")

            with gr.Column(scale=2):
                export_format = gr.Dropdown(
                    choices=["CSV", "PDF", "Word", "图片(A4)"],
                    label="导出格式",
                    value="PDF"
                )

            with gr.Column(scale=3):
                with gr.Row():
                    export_questions_button = gr.Button("导出题目", variant="secondary")
                    export_answers_button = gr.Button("导出答案", variant="secondary")

            with gr.Column(scale=2):
                delete_button = gr.Button("删除选中错题", variant="stop")

        # 添加文件下载组件
        download_file = gr.File(label="下载文件", visible=False, interactive=True)

        # 操作状态
        status = gr.Label(label="状态", value="")

        # 错题详情展示区
        with gr.Group(visible=False) as question_detail_group:
            gr.Markdown("## 错题详情")

            # 显示模式状态
            edit_mode = gr.State(False)

            # 添加基本操作按钮
            with gr.Row():
                back_button = gr.Button("返回列表", variant="secondary")
                edit_toggle_button = gr.Button("编辑错题", variant="primary")
                save_button = gr.Button("保存修改", variant="primary", visible=False)

            with gr.Tabs() as detail_tabs:
                with gr.TabItem("题目信息"):
                    # 只读模式的题目详情
                    with gr.Group(visible=True) as view_mode_group:
                        detail_markdown = gr.Markdown()

                        with gr.Row():
                            with gr.Column(scale=1):
                                subject_display = gr.Textbox(label="学科", interactive=False)
                                difficulty_display = gr.Textbox(label="难度", interactive=False)

                            with gr.Column(scale=1):
                                question_type_display = gr.Textbox(label="题型", interactive=False)
                                created_at_display = gr.Textbox(label="创建时间", interactive=False)

                        with gr.Row():
                            correct_answer = gr.Textbox(label="正确答案", interactive=False, lines=2)
                            user_answer = gr.Textbox(label="学生答案", interactive=False, lines=2)

                    # 编辑模式的题目详情
                    with gr.Group(visible=False) as edit_mode_group:
                        # 可编辑的表单
                        question_text_edit = gr.Textbox(label="题目内容", lines=5)

                        with gr.Row():
                            with gr.Column(scale=1):
                                subject_edit = gr.Dropdown(
                                    choices=["数学", "物理", "化学"],
                                    label="学科"
                                )
                                difficulty_edit = gr.Slider(
                                    minimum=1,
                                    maximum=5,
                                    step=1,
                                    label="难度等级",
                                    value=3
                                )

                            with gr.Column(scale=1):
                                question_type_edit = gr.Dropdown(
                                    choices=["选择题", "填空题", "判断题", "问答题"],
                                    label="题型"
                                )
                                # 创建时间不可编辑
                                created_at_edit = gr.Textbox(label="创建时间", interactive=False)

                        with gr.Row():
                            correct_answer_edit = gr.Textbox(label="正确答案", lines=2)
                            user_answer_edit = gr.Textbox(label="学生答案", lines=2)

                with gr.TabItem("解析"):
                    # 只读模式的解析
                    explanation = gr.Textbox(label="题目解析", lines=8, interactive=False)

                    # 编辑模式的解析
                    explanation_edit = gr.Textbox(label="题目解析", lines=8, visible=False)

        # 记录当前选中的行
        selected_row_index = gr.State(-1)
        selected_question_id = gr.State(None)

        # 加载所有错题数据
        def load_all_data():
            """加载所有错题数据"""
            try:
                # 验证数据一致性
                db.verify_data_consistency()

                questions = db.get_all_error_questions(limit=100)
                
                if not questions:
                    empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                    return "数据库中没有错题记录", empty_df
                
                # 验证每个问题的存在性并立即清理无效数据
                valid_questions = []
                invalid_ids = []
                
                for q in questions:
                    question = db.get_error_question(q["id"])
                    if question:
                        valid_questions.append(question)  # 使用完整的问题数据
                    else:
                        invalid_ids.append(q["id"])
                        # 立即从数据库中删除无效记录
                        db.delete_error_question(q["id"])
                        logger.warning(f"清理无效的错题ID: {q['id']}")
                
                if invalid_ids:
                    logger.warning(f"已清理 {len(invalid_ids)} 条无效错题记录: {invalid_ids}")
                
                if not valid_questions:
                    empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                    return "没有有效的错题记录", empty_df

                # 使用验证过的数据创建DataFrame
                data = {
                    "题目ID": [q["id"] for q in valid_questions],
                    "题目内容": [q["question_text"][:50] + "..." if len(q["question_text"]) > 50 else q["question_text"] for q in valid_questions],
                    "科目": [q["subject"] for q in valid_questions],
                    "题型": [q["question_type"] for q in valid_questions],
                    "难度": [f"{'⭐' * q['difficulty']}" for q in valid_questions],
                    "提交时间": [q["created_at"] for q in valid_questions]
                }
                df = pd.DataFrame(data)
                
                status_message = f"显示所有错题，共 {len(valid_questions)} 条"
                if invalid_ids:
                    status_message += f" (已清理 {len(invalid_ids)} 条无效记录)"
                
                return status_message, df
                
            except Exception as e:
                logger.error(f"加载数据错误: {str(e)}")
                logger.error(traceback.format_exc())
                empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                return f"加载数据出错: {str(e)}", empty_df

        def refresh_data():
            """刷新数据并更新界面"""
            try:
                # 执行数据一致性检查
                db.verify_data_consistency()
                
                # 获取最新数据
                status_msg, new_data = load_all_data()
                
                # 更新全局状态
                if isinstance(results, gr.DataFrame):
                    results.value = new_data
                if isinstance(query_status, gr.Textbox):
                    query_status.value = status_msg
                
                return status_msg, new_data
            except Exception as e:
                logger.error(f"刷新数据失败: {str(e)}")
                logger.error(traceback.format_exc())
                empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                error_msg = f"刷新数据失败: {str(e)}"
                return error_msg, empty_df

        def handle_refresh_click():
            """处理刷新按钮点击事件"""
            try:
                status_msg, new_data = refresh_data()
                # 清空详情区域
                return [
                    status_msg,  # 状态消息
                    new_data,    # 数据表格
                    gr.update(visible=True),  # 详情容器
                    False,       # 编辑状态
                    "### 请从列表中选择一道题目查看详情",  # 详情文本
                    *[""] * 7,   # 清空其他详情字段
                    -1,         # 重置行索引
                    None,       # 重置题目ID
                    gr.update(visible=True),   # 编辑按钮
                    gr.update(visible=False),  # 保存按钮
                    gr.update(visible=True),   # 删除按钮
                    gr.update(visible=False)   # 取消按钮
                ]
            except Exception as e:
                logger.error(f"刷新按钮处理失败: {str(e)}")
                logger.error(traceback.format_exc())
                empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                error_msg = f"刷新失败: {str(e)}"
                return [error_msg, empty_df] + [gr.update(value="") for _ in range(12)]

        # 搜索功能实现
        def search_questions(search_text, subjects, time_range):
            logger.info(f"搜索条件: {search_text}, {subjects}, {time_range}")
            try:
                # 更新查询状态
                query_status_text = "正在查询..."

                # 筛选条件
                subject_filter = subjects if subjects else None

                # 获取错题数据
                questions = db.get_all_error_questions(
                    subject=subject_filter,
                    time_range=time_range,
                    limit=100
                )

                # 如果有搜索关键词，进行过滤
                if search_text:
                    search_text = search_text.strip().lower()
                    filtered_questions = []

                    # 检查是否包含多个题号（用空格分隔）
                    search_terms = search_text.split()

                    # 如果全部是数字，则认为是搜索题号
                    are_all_numbers = all(term.isdigit() for term in search_terms)

                    for q in questions:
                        question_id_str = str(q["id"])

                        # 匹配多个题号的情况
                        if are_all_numbers:
                            # 精确匹配完整题号 - 检查题号是否在搜索词列表中
                            if question_id_str in search_terms:
                                filtered_questions.append(q)
                                continue

                            # 匹配题号中包含任一搜索数字
                            for term in search_terms:
                                if term in question_id_str:
                                    filtered_questions.append(q)
                                    break
                        # 按单个搜索词处理
                        else:
                            # 精确匹配题号
                            if question_id_str == search_text:
                                filtered_questions.append(q)
                                continue

                            # 题号包含搜索文本
                            if search_text in question_id_str:
                                filtered_questions.append(q)
                                continue

                            # 按题目内容搜索
                            if search_text in q["question_text"].lower():
                                filtered_questions.append(q)
                                continue

                    questions = filtered_questions

                # 转换为DataFrame显示格式
                if questions:
                    data = {
                        "题目ID": [q["id"] for q in questions],
                        "题目内容": [q["question_text"][:50] + "..." if len(q["question_text"]) > 50 else q["question_text"] for q in questions],
                        "科目": [q["subject"] for q in questions],
                        "题型": [q["question_type"] for q in questions],
                        "难度": [f"{'⭐' * q['difficulty']}" for q in questions],
                        "提交时间": [q["created_at"] for q in questions]
                    }
                    df = pd.DataFrame(data)
                    query_status_text = f"找到 {len(questions)} 条结果"
                else:
                    # 返回空DataFrame
                    df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                    query_status_text = "未找到符合条件的错题"

                return query_status_text, df

            except Exception as e:
                # 记录错误日志
                logger.error(f"搜索出错: {str(e)}")
                logger.error(traceback.format_exc())

                # 返回错误信息
                empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                query_status_text = f"搜索出错: {str(e)}"

                return query_status_text, empty_df

        # 处理DataFrame结构
        def process_dataframe_structure(df_value):
            """处理并转换非标准的DataFrame结构"""
            try:
                # 如果已经是标准DataFrame，检查是否含有所需列
                if isinstance(df_value, pd.DataFrame):
                    if "题目ID" in df_value.columns:
                        return df_value

                    # 输出列信息便于调试
                    logger.info(f"DataFrame列: {df_value.columns.tolist()}")

                    # 特殊情况：['headers', 'data', 'metadata']结构
                    if all(col in df_value.columns for col in ['headers', 'data']):
                        try:
                            # 尝试解析这种特殊结构
                            headers = df_value['headers'].tolist() if isinstance(df_value['headers'], pd.Series) else df_value['headers']
                            data = df_value['data'].tolist() if isinstance(df_value['data'], pd.Series) else df_value['data']

                            # 记录这些数据的结构
                            logger.info(f"Headers type: {type(headers)}, Data type: {type(data)}")

                            # 如果headers和data都是列表，可以尝试重构
                            if isinstance(headers, list) and isinstance(data, list):
                                # 确保headers中包含"题目ID"
                                if "题目ID" not in headers:
                                    logger.error("Headers中不包含'题目ID'列")
                                    return pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])

                                # 构建新的DataFrame
                                # 假设data是一个行列表，每一行是与headers对应的值列表
                                new_data = {}
                                for i, header in enumerate(headers):
                                    # 提取每一列的数据
                                    column_data = []
                                    for row in data:
                                        if i < len(row):
                                            column_data.append(row[i])
                                        else:
                                            column_data.append(None)
                                    new_data[header] = column_data

                                return pd.DataFrame(new_data)
                        except Exception as e:
                            logger.error(f"处理特殊DataFrame结构失败: {str(e)}")
                            logger.error(traceback.format_exc())

                # 处理字典类型
                elif isinstance(df_value, dict):
                    # 记录字典结构
                    for k, v in df_value.items():
                        logger.info(f"Key: {k}, Type: {type(v)}, Length: {len(v) if hasattr(v, '__len__') else 'N/A'}")

                    # 处理特殊字典结构
                    if 'headers' in df_value and 'data' in df_value:
                        headers = df_value['headers']
                        data = df_value['data']

                        if isinstance(headers, list) and isinstance(data, list):
                            # 同上，构建新的字典结构
                            new_data = {}
                            for i, header in enumerate(headers):
                                column_data = []
                                for row in data:
                                    if i < len(row):
                                        column_data.append(row[i])
                                    else:
                                        column_data.append(None)
                                new_data[header] = column_data

                            return pd.DataFrame(new_data)

                    # 尝试规范化普通字典
                    try:
                        normalized_dict = normalize_dict(df_value)
                        return pd.DataFrame(normalized_dict)
                    except Exception as e:
                        logger.error(f"规范化字典失败: {str(e)}")

                # 其他情况，返回空DataFrame
                logger.error(f"无法处理的数据类型: {type(df_value)}")
                return pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])

            except Exception as e:
                logger.error(f"处理DataFrame结构时出错: {str(e)}")
                logger.error(traceback.format_exc())
                return pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])

        # 规范化字典函数，确保所有键的值具有相同的长度
        def normalize_dict(d):
            # 找出字典中非空列表值的最大长度
            max_len = 0
            for v in d.values():
                if isinstance(v, list) and len(v) > max_len:
                    max_len = len(v)

            # 规范化所有值为相同长度的列表
            result = {}
            for k, v in d.items():
                if isinstance(v, list):
                    # 如果是列表但长度不够，填充None
                    if len(v) < max_len:
                        result[k] = v + [None] * (max_len - len(v))
                    else:
                        result[k] = v
                else:
                    # 如果不是列表，创建一个填充了相同值的列表
                    result[k] = [v] * max_len if max_len > 0 else []

            return result

        # 切换编辑模式
        def toggle_edit_mode(edit_state, question_id):
            if not question_id:
                return edit_state, gr.update(visible=True), gr.update(visible=False), gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), gr.update(visible=False)

            # 切换编辑状态
            new_state = not edit_state

            if new_state:  # 进入编辑模式
                # 查询当前题目详情
                question = db.get_error_question(question_id)
                if not question:
                    logger.warning(f"编辑不存在的错题: ID={question_id}")
                    return edit_state, gr.update(visible=True), gr.update(visible=False), gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), gr.update(visible=False)

                # 填充编辑表单数据
                return (
                    new_state,                          # 新的编辑状态
                    gr.update(visible=False),           # 隐藏查看模式组
                    gr.update(visible=True),            # 显示编辑模式组
                    gr.update(visible=False),           # 隐藏解析（只读模式）
                    gr.update(visible=True, value=str(question.get("explanation", ""))),  # 显示解析编辑框，确保是字符串
                    gr.update(visible=False),           # 隐藏编辑按钮
                    gr.update(visible=True),            # 显示保存按钮
                    str(question.get("question_text", "")),          # 题目内容
                    str(question.get("subject", "")),                # 学科
                    str(question.get("question_type", "")),          # 题型
                    int(question.get("difficulty", 3)),             # 难度
                    str(question.get("created_at", "")),             # 创建时间
                    str(question.get("answer", "")),                 # 正确答案
                    str(question.get("user_answer", ""))             # 用户答案
                )
            else:  # 退出编辑模式
                return (
                    new_state,                     # 新的编辑状态
                    gr.update(visible=True),       # 显示查看模式组
                    gr.update(visible=False),      # 隐藏编辑模式组
                    gr.update(visible=True),       # 显示解析（只读模式）
                    gr.update(visible=False),      # 隐藏解析编辑框
                    gr.update(visible=True),       # 显示编辑按钮
                    gr.update(visible=False),      # 隐藏保存按钮
                )

        # 保存编辑后的错题
        def save_edited_question(
            question_id,
            question_text,
            subject,
            question_type,
            difficulty,
            correct_answer,
            user_answer,
            explanation
        ):
            try:
                if not question_id:
                    return "❌ 保存失败：未选择题目", False

                # 验证输入数据并确保类型正确
                if not str(question_text).strip():
                    return "❌ 保存失败：题目内容不能为空", False

                if not subject:
                    return "❌ 保存失败：请选择学科", False

                if not question_type:
                    return "❌ 保存失败：请选择题型", False
                    
                # 确保难度是整数并在有效范围内
                try:
                    difficulty = int(difficulty)
                    if not (1 <= difficulty <= 5):
                        difficulty = max(1, min(5, difficulty))  # 限制在1-5范围内
                except (TypeError, ValueError):
                    difficulty = 3  # 默认值
                    
                # 构建更新数据
                update_data = {
                    "question_text": str(question_text),
                    "subject": str(subject),
                    "question_type": str(question_type),
                    "difficulty": int(difficulty),
                    "answer": str(correct_answer or ""),
                    "user_answer": str(user_answer or ""),
                    "explanation": str(explanation or "")
                }

                # 更新数据库
                success = db.update_error_question(question_id, update_data)

                if success:
                    # 只返回状态消息和成功标志
                    return "✅ 保存成功！", True
                else:
                    return "❌ 保存失败：可能是数据库错误", False

            except Exception as e:
                logger.error(f"保存编辑错题时出错: {str(e)}")
                logger.error(traceback.format_exc())
                return f"❌ 保存失败: {str(e)}", False

        # 显示错题详情
        # 在paste-2.txt文件中找到show_question_detail函数
        def show_question_detail(evt: gr.SelectData):
            """显示错题详情"""
            try:
                row_index = evt.index[0] if evt and hasattr(evt, 'index') else -1
                df_value = results.value if results else None
                
                # 使用process_dataframe_structure处理df_value
                if df_value is not None:
                    try:
                        df_value = process_dataframe_structure(df_value)
                        logger.info(f"处理后的DataFrame列: {df_value.columns.tolist() if hasattr(df_value, 'columns') else 'N/A'}")
                    except Exception as e:
                        logger.error(f"处理DataFrame结构失败: {str(e)}")
                        return False, f"处理数据格式失败: {str(e)}", "", "", "", "", "", "", "", -1, None, "", gr.update(visible=True), gr.update(visible=False)
                
                # 检查df_value的有效性
                if df_value is None or not isinstance(df_value, pd.DataFrame) or df_value.empty or "题目ID" not in df_value.columns:
                    logger.warning("数据无效或缺少题目ID列，正在刷新...")
                    return False, "数据无效，请刷新页面", "", "", "", "", "", "", "", -1, None, "", gr.update(visible=True), gr.update(visible=False)
                
                # 检查行索引是否有效
                if not (0 <= row_index < len(df_value)):
                    logger.warning(f"行索引{row_index}超出范围(0-{len(df_value)-1})")
                    return False, f"选择的行索引{row_index}无效，请重新选择", "", "", "", "", "", "", "", -1, None, "", gr.update(visible=True), gr.update(visible=False)

                # 获取题目ID
                try:
                    question_id = int(df_value.iloc[row_index]["题目ID"])
                    logger.info(f"成功获取题目ID: {question_id}")
                except (ValueError, TypeError, IndexError) as e:
                    logger.error(f"获取题目ID失败: {str(e)}")
                    return False, f"获取题目ID失败: {str(e)}", "", "", "", "", "", "", "", -1, None, "", gr.update(visible=True), gr.update(visible=False)

                # 查询题目详情
                question = db.get_error_question(question_id)
                
                # 如果题目不存在
                if not question:
                    logger.warning(f"题目ID={question_id}不存在")
                    error_message = f"### ⚠️ 错题不存在\n\n该错题（ID: {question_id}）可能已被删除或ID无效。"
                    return False, error_message, "", "", "", "", "", "", "", -1, None, "", gr.update(visible=True), gr.update(visible=False)

                # 题目存在，显示详情
                logger.info(f"成功找到题目: ID={question_id}")
                difficulty_stars = "⭐" * question["difficulty"]
                detail_text = f"""### 题目内容\n\n{question["question_text"]}\n\n### 基本信息\n- 科目：{question["subject"]}\n- 题型：{question["question_type"]}\n- 难度：{difficulty_stars}\n- 提交时间：{question["created_at"]}"""
                
                # 确保所有值都是字符串类型
                subject = str(question.get("subject", ""))
                question_type = str(question.get("question_type", ""))
                created_at = str(question.get("created_at", ""))
                answer = str(question.get("answer", ""))
                user_answer = str(question.get("user_answer", ""))
                explanation = str(question.get("explanation", "无解析"))
                
                # 关键修改：不要直接设置组件的visible属性，而是返回gr.update()
                # 删除这些直接设置可见性的代码:
                # question_detail_group.visible = True
                # view_mode_group.visible = True
                # edit_mode_group.visible = False
                
                # 返回值 - 确保每一个值都是正确的类型，且与outputs列表匹配
                return False, detail_text, subject, question_type, difficulty_stars, created_at, answer, user_answer, explanation, row_index, question_id, explanation, gr.update(visible=True), gr.update(visible=False)

            except Exception as e:
                logger.error(f"显示错题详情时出错: {str(e)}")
                logger.error(traceback.format_exc())
                # 返回安全的默认值
                return False, f"显示详情时出错: {str(e)}", "", "", "", "", "", "", "", -1, None, "", gr.update(visible=True), gr.update(visible=False)
        # 添加规范化字典的函数
        def normalize_dict(d):
            """确保字典中所有键的值具有相同的长度"""
            # 找出字典中非空列表值的最大长度
            max_len = 0
            for v in d.values():
                if isinstance(v, list) and len(v) > max_len:
                    max_len = len(v)

            # 规范化所有值为相同长度的列表
            result = {}
            for k, v in d.items():
                if isinstance(v, list):
                    # 如果是列表但长度不够，填充None
                    if len(v) < max_len:
                        result[k] = v + [None] * (max_len - len(v))
                    else:
                        result[k] = v
                else:
                    # 如果不是列表，创建一个填充了相同值的列表
                    result[k] = [v] * max_len if max_len > 0 else []

            return result
        # 辅助函数：文本换行处理
        def textwrap_text(text, font, max_width):
            """将文本按照给定宽度分割成多行"""
            words = text.split()
            lines = []
            current_line = words[0] if words else ""

            for word in words[1:]:
                # 计算添加新单词后的宽度
                test_line = current_line + " " + word
                width = font.getmask(test_line).getbbox()[2]

                if width <= max_width:
                    current_line = test_line
                else:
                    lines.append(current_line)
                    current_line = word

            lines.append(current_line)  # 添加最后一行
            return lines

        # 修改导出功能 - 支持直接下载
        def export_results(dataframe, export_format="CSV", content_type="questions"):
            """
            导出错题数据

            参数:
            dataframe: 要导出的数据
            export_format: 导出格式，可选 CSV, PDF, Word, 图片(A4)
            content_type: 内容类型，可选 questions(只有题目), answers(只有答案)
            """
            if dataframe is None:
                return "没有可导出的数据", None, gr.update(visible=False)

            try:
                # 处理数据结构
                df = process_dataframe_structure(dataframe)

                if df.empty:
                    return "没有可导出的数据", None, gr.update(visible=False)

                # 获取所有ID对应的完整数据
                if "题目ID" not in df.columns:
                    return f"❌ 导出失败: 数据缺少'题目ID'列", None, gr.update(visible=False)

                ids = df["题目ID"].tolist()
                full_data = []
                missing_ids = []

                for qid in ids:
                    if pd.isna(qid):  # 跳过NaN值
                        continue
                    question = db.get_error_question(qid)
                    if question:
                        full_data.append(question)
                    else:
                        # 记录不存在的题目ID
                        missing_ids.append(str(qid))

                # 检查是否有缺失的题目
                if missing_ids:
                    logger.warning(f"导出时发现不存在的题目ID: {', '.join(missing_ids)}")
                    if not full_data:  # 如果没有有效数据
                        return f"❌ 导出失败: 所选题目均不存在 (缺失ID: {', '.join(missing_ids)})", None, gr.update(visible=False)

                # 创建导出目录
                export_dir = os.path.join(os.getcwd(), "exports")
                os.makedirs(export_dir, exist_ok=True)
                timestamp = time.strftime("%Y%m%d_%H%M%S")

                # 根据内容类型设置文件名
                if content_type == "questions":
                    file_suffix = "题目"
                else:
                    file_suffix = "答案"

                # 根据选择的格式导出
                if export_format == "CSV":
                    # CSV导出
                    csv_path = os.path.join(export_dir, f"错题_{file_suffix}_{timestamp}.csv")

                    with open(csv_path, 'w', newline='', encoding='utf-8-sig') as f:
                        if content_type == "questions":
                            # 只导出题目部分
                            fieldnames = ["id", "subject", "question_type", "difficulty", "question_text"]
                        else:
                            # 只导出答案部分
                            fieldnames = ["id", "subject", "question_type", "answer", "explanation"]

                        writer = csv.DictWriter(f, fieldnames=fieldnames)
                        writer.writeheader()

                        for item in full_data:
                            # 筛选需要的字段
                            row = {field: item.get(field, '') for field in fieldnames}
                            writer.writerow(row)

                    success_msg = f"✅ 已成功导出 {len(full_data)} 条{file_suffix}至CSV文件，请点击下载"
                    if missing_ids:
                        success_msg += f" (注意: {len(missing_ids)} 条题目不存在，已跳过ID: {', '.join(missing_ids[:3])}{', ...' if len(missing_ids) > 3 else ''})"
                    return success_msg, csv_path, gr.update(visible=True)

                elif export_format == "PDF":
                    try:
                        # 导入所需模块
                        import reportlab.lib.pagesizes as pagesizes
                        import reportlab.platypus as platypus
                        import reportlab.lib.styles as styles
                        import reportlab.lib.colors as colors
                        import reportlab.lib.enums as enums

                        # 添加中文字体支持
                        from reportlab.pdfbase import pdfmetrics
                        from reportlab.pdfbase.ttfonts import TTFont

                        # 注册中文字体 - 尝试多个路径
                        font_registered = False
                        font_paths = [
                            # Windows 字体路径
                            "C:/Windows/Fonts/simhei.ttf",
                            "C:/Windows/Fonts/simsun.ttc",
                            # Linux 字体路径
                            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                            # macOS 字体路径
                            "/System/Library/Fonts/PingFang.ttc",
                            # 项目内字体（如果有）
                            os.path.join(os.getcwd(), "fonts", "simhei.ttf"),
                        ]

                        for font_path in font_paths:
                            if os.path.exists(font_path):
                                try:
                                    logger.info(f"尝试注册字体: {font_path}")
                                    # 注册字体
                                    pdfmetrics.registerFont(TTFont('SimHei', font_path))
                                    font_registered = True
                                    logger.info(f"成功注册字体: {font_path}")
                                    break
                                except Exception as font_error:
                                    logger.error(f"注册字体失败: {str(font_error)}")
                                    continue

                        # 创建PDF文档
                        pdf_path = os.path.join(export_dir, f"错题_{file_suffix}_{timestamp}.pdf")
                        doc = platypus.SimpleDocTemplate(pdf_path, pagesize=pagesizes.A4)

                        # 创建自定义样式
                        style_sheet = styles.getSampleStyleSheet()

                        # 修改所有样式以使用中文字体
                        if font_registered:
                            for style_name in style_sheet.byName:
                                style_sheet[style_name].fontName = 'SimHei'

                        title_style = styles.ParagraphStyle(
                            'CustomTitle',
                            parent=style_sheet['Title'],
                            fontName='SimHei' if font_registered else style_sheet['Title'].fontName,
                            textColor=colors.HexColor('#003399'),
                            alignment=enums.TA_CENTER,
                            spaceAfter=10
                        )

                        heading_style = styles.ParagraphStyle(
                            'CustomHeading',
                            parent=style_sheet['Heading2'],
                            fontName='SimHei' if font_registered else style_sheet['Heading2'].fontName,
                            textColor=colors.HexColor('#0066CC'),
                            spaceBefore=15
                        )

                        elements = []

                        # 添加标题
                        elements.append(platypus.Paragraph(f"错题集 - {file_suffix}", title_style))
                        elements.append(platypus.Spacer(1, 12))

                        # 遍历每道题目
                        for i, question in enumerate(full_data, 1):
                            if content_type == "questions":
                                # 导出题目部分
                                # 题目标题
                                elements.append(platypus.Paragraph(f"题目 {i}: {question['subject']} - {question['question_type']}", heading_style))

                                # 题目内容
                                elements.append(platypus.Paragraph(question["question_text"], style_sheet['Normal']))
                            else:
                                # 导出答案部分
                                # 标题简化显示题号
                                elements.append(platypus.Paragraph(f"题目 {i} 答案:", heading_style))

                                # 显示答案
                                elements.append(platypus.Paragraph(f"<b><font color='#008000'>正确答案:</font></b> {question['answer']}", style_sheet['Normal']))

                                # 解析
                                if question.get("explanation"):
                                    elements.append(platypus.Spacer(1, 5))
                                    elements.append(platypus.Paragraph("<b><font color='#4682B4'>解析:</font></b>", style_sheet['Normal']))
                                    elements.append(platypus.Paragraph(question["explanation"], style_sheet['Normal']))

                            # 添加分隔线
                            if i < len(full_data):  # 最后一题不添加分隔线
                                elements.append(platypus.Spacer(1, 10))

                                # 创建简单表格作为分隔线
                                data = [['']]
                                line_table = platypus.Table(data, colWidths=[450], rowHeights=[1])
                                line_table.setStyle([('LINEABOVE', (0, 0), (-1, -1), 1, colors.gray)])
                                elements.append(line_table)

                                elements.append(platypus.Spacer(1, 15))

                        # 生成PDF
                        doc.build(elements)
                        success_msg = f"✅ 已成功导出 {len(full_data)} 条错题{file_suffix}至PDF文件，请点击下载"
                        if missing_ids:
                            success_msg += f" (注意: {len(missing_ids)} 条题目不存在，已跳过ID: {', '.join(missing_ids[:3])}{', ...' if len(missing_ids) > 3 else ''})"
                        return success_msg, pdf_path, gr.update(visible=True)

                    except Exception as e:
                        logger.error(f"PDF导出过程中出错: {str(e)}")
                        logger.error(traceback.format_exc())

                        # 检查是否是ImportError
                        if isinstance(e, ImportError):
                            return f"❌ 导出PDF需要安装reportlab库，请使用命令：pip install reportlab", None, gr.update(visible=False)
                        else:
                            return f"❌ PDF导出失败: {str(e)}", None, gr.update(visible=False)

                elif export_format == "Word":
                    try:
                        # 使用python-docx导出Word文档
                        from docx import Document
                        from docx.shared import RGBColor

                        doc_path = os.path.join(export_dir, f"错题_{file_suffix}_{timestamp}.docx")
                        doc = Document()

                        # 添加标题
                        title = doc.add_heading(f"错题集 - {file_suffix}", 0)
                        title.alignment = 1  # 居中对齐

                        # 遍历每道题目
                        for i, question in enumerate(full_data, 1):
                            if content_type == "questions":
                                # 导出题目部分
                                # 题目标题
                                heading = doc.add_heading(f"题目 {i}: {question['subject']} - {question['question_type']}", level=2)

                                # 设置标题颜色为蓝色
                                for run in heading.runs:
                                    run.font.color.rgb = RGBColor(0, 102, 204)

                                # 题目内容
                                doc.add_paragraph(question["question_text"])
                            else:
                                # 导出答案部分
                                # 标题
                                heading = doc.add_heading(f"题目 {i} 答案:", level=2)

                                # 设置标题颜色为蓝色
                                for run in heading.runs:
                                    run.font.color.rgb = RGBColor(0, 102, 204)

                                # 答案
                                p = doc.add_paragraph()
                                correct_run = p.add_run("正确答案: ")
                                correct_run.bold = True
                                correct_run.font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                                p.add_run(question['answer'])

                                # 解析
                                if question.get("explanation"):
                                    p = doc.add_paragraph()
                                    explain_run = p.add_run("解析: ")
                                    explain_run.bold = True
                                    explain_run.font.color.rgb = RGBColor(70, 130, 180)  # 钢蓝色
                                    p.add_run(question["explanation"])

                            # 除最后一题外，添加分隔线
                            if i < len(full_data):
                                doc.add_paragraph("_" * 50)

                        # 保存文档
                        doc.save(doc_path)
                        success_msg = f"✅ 已成功导出 {len(full_data)} 条错题{file_suffix}至Word文档，请点击下载"
                        if missing_ids:
                            success_msg += f" (注意: {len(missing_ids)} 条题目不存在，已跳过ID: {', '.join(missing_ids[:3])}{', ...' if len(missing_ids) > 3 else ''})"
                        return success_msg, doc_path, gr.update(visible=True)
                    except ImportError:
                        return "❌ 导出Word需要安装python-docx库，请使用命令：pip install python-docx", None, gr.update(visible=False)
                    except Exception as e:
                        logger.error(f"Word导出失败: {str(e)}")
                        logger.error(traceback.format_exc())
                        return f"❌ Word导出失败: {str(e)}", None, gr.update(visible=False)

                elif export_format == "图片(A4)":
                    try:
                        # 使用PIL生成A4尺寸的可打印图片
                        from PIL import Image, ImageDraw, ImageFont
                        import math
                        import zipfile

                        # 创建图片保存目录
                        img_dir = os.path.join(export_dir, f"错题_{file_suffix}_{timestamp}")
                        os.makedirs(img_dir, exist_ok=True)

                        # A4纸像素尺寸（300DPI）
                        width, height = 2480, 3508  # A4尺寸，300DPI

                        # 计算需要多少页
                        items_per_page = 3  # 每页题目数
                        total_pages = math.ceil(len(full_data) / items_per_page)

                        # 字体设置 - 使用更通用的字体路径或内置字体
                        # 尝试多种字体路径以提高兼容性
                        font_paths = [
                            # Windows 字体路径
                            "C:/Windows/Fonts/simhei.ttf",
                            "C:/Windows/Fonts/simsun.ttc",
                            # Linux 字体路径
                            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                            # macOS 字体路径
                            "/System/Library/Fonts/PingFang.ttc",
                            # 项目内字体（如果有）
                            os.path.join(os.getcwd(), "fonts", "simhei.ttf"),
                        ]

                        # 尝试加载字体，直到成功或全部失败
                        title_font = None
                        heading_font = None
                        normal_font = None

                        for font_path in font_paths:
                            if os.path.exists(font_path):
                                try:
                                    title_font = ImageFont.truetype(font_path, 60)
                                    heading_font = ImageFont.truetype(font_path, 45)
                                    normal_font = ImageFont.truetype(font_path, 35)
                                    logger.info(f"成功加载字体: {font_path}")
                                    break
                                except Exception as e:
                                    logger.error(f"加载字体 {font_path} 失败: {e}")
                                    continue

                        # 如果所有字体都加载失败，使用默认字体
                        if title_font is None:
                            logger.warning("无法加载任何中文字体，使用默认字体")
                            title_font = ImageFont.load_default()
                            heading_font = ImageFont.load_default()
                            normal_font = ImageFont.load_default()

                        # 生成的图片路径列表
                        image_paths = []

                        # 生成每一页
                        for page in range(total_pages):
                            # 创建空白图片（白色背景）
                            img = Image.new('RGB', (width, height), (255, 255, 255))
                            draw = ImageDraw.Draw(img)

                            # 绘制标题和页眉线条
                            draw.text((width//2 - 250, 100), f"错题集 - {file_suffix}", font=title_font, fill=(0, 51, 153))  # 深蓝色标题
                            draw.line([(80, 170), (width-80, 170)], fill=(0, 102, 204), width=2)  # 蓝色分隔线

                            # 当前页的题目
                            start_idx = page * items_per_page
                            end_idx = min(start_idx + items_per_page, len(full_data))
                            current_questions = full_data[start_idx:end_idx]

                            # 绘制每道题目
                            y_position = 250
                            for i, question in enumerate(current_questions, start_idx + 1):
                                if content_type == "questions":
                                    # 题目部分
                                    # 题目标题 - 使用蓝色
                                    title_text = f"题目 {i}: {question['subject']} - {question['question_type']}"
                                    draw.text((100, y_position), title_text, font=heading_font, fill=(0, 102, 204))
                                    y_position += 70

                                    # 题目内容
                                    content_text = question["question_text"]

                                    # 手动处理文本换行 - 简化处理
                                    # 每行最大字符数（根据字体大小估算）
                                    max_chars_per_line = 65

                                    # 分行处理
                                    lines = []
                                    for paragraph in content_text.split('\n'):
                                        if len(paragraph) <= max_chars_per_line:
                                            lines.append(paragraph)
                                        else:
                                            # 长段落按字符数分行
                                            for j in range(0, len(paragraph), max_chars_per_line):
                                                lines.append(paragraph[j:j+max_chars_per_line])

                                    # 绘制题目内容
                                    for line in lines:
                                        draw.text((100, y_position), line, font=normal_font, fill=(0, 0, 0))
                                        y_position += 45
                                else:
                                    # 答案部分
                                    # 标题
                                    title_text = f"题目 {i} 答案:"
                                    draw.text((100, y_position), title_text, font=heading_font, fill=(0, 102, 204))
                                    y_position += 70

                                    # 答案
                                    draw.text((100, y_position), "正确答案:", font=normal_font, fill=(0, 128, 0))  # 绿色
                                    draw.text((300, y_position), question['answer'], font=normal_font, fill=(0, 0, 0))
                                    y_position += 70

                                    # 解析
                                    if question.get("explanation"):
                                        draw.text((100, y_position), "解析:", font=normal_font, fill=(70, 130, 180))  # 钢蓝色
                                        y_position += 50

                                        # 分行处理解析文本
                                        explanation_lines = []
                                        max_chars = 65  # 每行最大字符数
                                        for paragraph in question["explanation"].split('\n'):
                                            if len(paragraph) <= max_chars:
                                                explanation_lines.append(paragraph)
                                            else:
                                                for j in range(0, len(paragraph), max_chars):
                                                    explanation_lines.append(paragraph[j:j+max_chars])

                                        for line in explanation_lines:
                                            draw.text((100, y_position), line, font=normal_font, fill=(0, 0, 0))
                                            y_position += 45

                                # 添加分隔线
                                y_position += 50
                                draw.line([(150, y_position), (width-150, y_position)], fill=(200, 200, 200), width=1)
                                y_position += 60

                            # 添加页码和页脚线
                            draw.line([(80, height-150), (width-80, height-150)], fill=(0, 102, 204), width=1)
                            draw.text((width//2 - 80, height - 100), f"第 {page+1}/{total_pages} 页",
                                    font=normal_font, fill=(102, 102, 102))

                            # 保存图片
                            img_path = os.path.join(img_dir, f"错题_{file_suffix}_第{page+1}页.png")
                            img.save(img_path, "PNG")
                            image_paths.append(img_path)

                        # 创建ZIP文件以供下载
                        zip_path = os.path.join(export_dir, f"错题_{file_suffix}_{timestamp}.zip")
                        with zipfile.ZipFile(zip_path, 'w') as zipf:
                            for img_path in image_paths:
                                # 仅添加文件名而非完整路径
                                zipf.write(img_path, os.path.basename(img_path))

                        success_msg = f"✅ 已成功导出 {len(full_data)} 条错题{file_suffix}至 {total_pages} 页图片，请点击下载"
                        if missing_ids:
                            success_msg += f" (注意: {len(missing_ids)} 条题目不存在，已跳过ID: {', '.join(missing_ids[:3])}{', ...' if len(missing_ids) > 3 else ''})"
                        return success_msg, zip_path, gr.update(visible=True)
                    except ImportError:
                        return "❌ 导出图片需要安装Pillow库，请使用命令：pip install Pillow", None, gr.update(visible=False)
                    except Exception as e:
                        logger.error(f"导出图片时出错: {str(e)}")
                        logger.error(traceback.format_exc())
                        return f"❌ 导出图片时出错: {str(e)}", None, gr.update(visible=False)

                else:
                    return f"❌ 不支持的导出格式: {export_format}", None, gr.update(visible=False)

            except Exception as e:
                logger.error(f"导出失败: {str(e)}")
                logger.error(traceback.format_exc())
                return f"❌ 导出失败: {str(e)}", None, gr.update(visible=False)

        # 分别导出题目和答案的函数
        def export_questions(dataframe, export_format):
            return export_results(dataframe, export_format, content_type="questions")

        def export_answers(dataframe, export_format):
            return export_results(dataframe, export_format, content_type="answers")

        # 返回列表视图
        def back_to_list():
            return [
                gr.update(visible=False),  # 修改：使用gr.update来控制Group的可见性
                False,                     # edit_mode
                "",                        # detail_markdown
                "",                        # subject_display
                "",                        # question_type_display
                "",                        # difficulty_display
                "",                        # created_at_display
                "",                        # correct_answer
                "",                        # user_answer
                "",                        # explanation
                -1,                        # selected_row_index
                None,                      # selected_question_id
                gr.update(visible=True),   # view_mode_group
                gr.update(visible=False),  # edit_mode_group
                gr.update(visible=True),   # explanation
                gr.update(visible=False),  # explanation_edit
                gr.update(visible=True),   # edit_toggle_button
                gr.update(visible=False)   # save_button
            ]

        # 直接删除错题函数
        # def direct_delete(question_id):
        #     if not question_id:
        #         return "请先选择要删除的错题", query_status.value, results.value, gr.update(visible=False)

        #     try:
        #         # 确保 question_id 是整数
        #         try:
        #             question_id = int(question_id) if question_id else None
        #         except (ValueError, TypeError):
        #             return f"❌ 删除失败：无效的错题ID {question_id}", query_status.value, results.value, gr.update(visible=False)

        #         # 在删除之前先检查错题是否存在
        #         question = db.get_error_question(question_id)
        #         if not question:
        #             return f"❌ 删除失败：未找到ID为 {question_id} 的错题", query_status.value, results.value, gr.update(visible=False)

        #         # 执行删除操作
        #         success = db.delete_error_question(question_id)

        #         # 获取最新数据
        #         refresh_msg, refreshed_data = load_all_data()

        #         if success:
        #             return f"✅ 成功删除错题(ID: {question_id})", refresh_msg, refreshed_data, gr.update(visible=False)
        #         else:
        #             return f"❌ 删除失败：操作未成功完成", refresh_msg, refreshed_data, gr.update(visible=False)

        #     except Exception as e:
        #         logger.error(f"删除错题时出错: {str(e)}")
        #         logger.error(traceback.format_exc())

        #         # 尝试刷新数据
        #         try:
        #             refresh_msg, refreshed_data = load_all_data()
        #             return f"❌ 删除错题时出错: {str(e)}", refresh_msg, refreshed_data, gr.update(visible=False)
        #         except:
        #             empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
        #             return f"❌ 删除错题时出错: {str(e)}", "加载数据失败", empty_df, gr.update(visible=False)
        def direct_delete(question_id):
            if not question_id:
                return "请先选择要删除的错题", query_status.value, results.value, gr.update(visible=False)  # 修改：使用gr.update控制可见性
            
            try:
                # 删除错题
                success = db.delete_error_question(question_id)
                
                # 获取最新数据
                refresh_msg, refreshed_data = load_all_data()
                
                if success:
                    status_msg = f"✅ 成功删除错题(ID: {question_id})"
                else:
                    status_msg = f"❌ 删除失败：未找到ID为 {question_id} 的错题"
                
                # 使用gr.update控制可见性
                return status_msg, refresh_msg, refreshed_data, gr.update(visible=False)
            except Exception as e:
                logger.error(f"删除错题时出错: {str(e)}")
                logger.error(traceback.format_exc())
                
                # 尝试刷新数据
                try:
                    refresh_msg, refreshed_data = load_all_data()
                    return f"❌ 删除错题时出错: {str(e)}", refresh_msg, refreshed_data, gr.update(visible=False)
                except:
                    empty_df = pd.DataFrame(columns=["题目ID", "题目内容", "科目", "题型", "难度", "提交时间"])
                    return f"❌ 删除错题时出错: {str(e)}", "加载数据失败", empty_df, gr.update(visible=False)
        # 保存编辑的处理函数
        def handle_save_edit(
            question_id,
            question_text,
            subject,
            question_type,
            difficulty,
            correct_answer,
            user_answer,
            explanation
        ):
            try:
                # 确保所有值都是合适的类型
                if question_id is None:
                    return "❌ 保存失败：未选择题目", True, gr.update(visible=False), gr.update(visible=True), results.value, gr.update(visible=False), gr.update(visible=True), gr.update(visible=False), gr.update(visible=True)
                    
                # 修改这里的调用方式，处理save_edited_question可能返回多个值的情况
                result = save_edited_question(
                    question_id,
                    str(question_text),
                    str(subject),
                    str(question_type),
                    int(difficulty),
                    str(correct_answer),
                    str(user_answer),
                    str(explanation)
                )
                
                # 检查返回值类型，并相应处理
                if isinstance(result, tuple) and len(result) >= 2:
                    # 如果返回多个值，第一个是状态消息，第二个是成功标志
                    status_msg = result[0]
                    success = result[1]
                    
                    # 如果成功且返回了完整的UI更新数据，直接使用
                    if success and len(result) > 2:
                        # 假设返回的是完整的UI更新列表
                        return result
                    elif success:
                        # 成功但只返回了状态和标志，需要刷新数据
                        refresh_msg, refreshed_data = load_all_data()
                        return [
                            status_msg,               # 状态消息
                            False,                    # 退出编辑模式
                            gr.update(visible=True),  # 显示查看模式
                            gr.update(visible=False), # 隐藏编辑模式
                            refreshed_data,           # 更新后的数据表格
                            gr.update(visible=True),  # 显示解析只读框
                            gr.update(visible=False), # 隐藏解析编辑框
                            gr.update(visible=True),  # 显示编辑按钮
                            gr.update(visible=False)  # 隐藏保存按钮
                        ]
                else:
                    # 如果只返回一个值，假设是错误消息
                    status_msg = str(result)
                    success = False
                    
                    return [
                        status_msg,
                        True,  # 保持编辑模式
                        gr.update(visible=False),
                        gr.update(visible=True),
                        results.value,
                        gr.update(visible=False),
                        gr.update(visible=True),
                        gr.update(visible=False),
                        gr.update(visible=True)
                    ]

            except Exception as e:
                logger.error(f"保存编辑失败: {str(e)}")
                logger.error(traceback.format_exc())
                return [
                    f"❌ 保存失败: {str(e)}",
                    True,  # 保持编辑模式
                    gr.update(visible=False),
                    gr.update(visible=True),
                    results.value,
                    gr.update(visible=False),
                    gr.update(visible=True),
                    gr.update(visible=False),
                    gr.update(visible=True)
                ]

        # 绑定事件
        search_button.click(
            search_questions,
            inputs=[search_input, filters, date_range],
            outputs=[query_status, results]
        )

        # 绑定表格选择事件
        results.select(
            show_question_detail,
            None,  # 输入参数为空，因为我们使用的是事件对象
            [
                edit_mode,
                detail_markdown,
                subject_display,
                question_type_display,
                difficulty_display,
                created_at_display,
                correct_answer,
                user_answer,
                explanation,
                selected_row_index,
                selected_question_id,
                explanation_edit,
                question_detail_group,  # 确保这里包含了question_detail_group
                edit_mode_group  # 确保这里包含了edit_mode_group
            ]
        )
        # 编辑模式切换
        edit_toggle_button.click(
            toggle_edit_mode,
            inputs=[edit_mode, selected_question_id],
            outputs=[
                edit_mode,
                view_mode_group,
                edit_mode_group,
                explanation,
                explanation_edit,
                edit_toggle_button,
                save_button,
                question_text_edit,
                subject_edit,
                question_type_edit,
                difficulty_edit,
                created_at_edit,
                correct_answer_edit,
                user_answer_edit
            ]
        )

        # 保存按钮点击事件
        save_button.click(
            handle_save_edit,
            inputs=[
                selected_question_id,
                question_text_edit,
                subject_edit,
                question_type_edit,
                difficulty_edit,
                correct_answer_edit,
                user_answer_edit,
                explanation_edit
            ],
            outputs=[
                status,
                edit_mode,
                view_mode_group,
                edit_mode_group,
                results,
                explanation,
                explanation_edit,
                edit_toggle_button,
                save_button
            ]
        )
        back_button.click(
            back_to_list,
            None,
            [
                question_detail_group,
                edit_mode,
                detail_markdown,
                subject_display,
                question_type_display,
                difficulty_display,
                created_at_display,
                correct_answer,
                user_answer,
                explanation,
                selected_row_index,
                selected_question_id,
                view_mode_group,
                edit_mode_group,
                explanation,
                explanation_edit,
                edit_toggle_button,
                save_button
            ]
        )

        # 修改导出题目按钮绑定 - 增加文件下载组件作为输出
        export_questions_button.click(
            export_questions,
            inputs=[results, export_format],
            outputs=[status, download_file, download_file]  # 第一个参数控制文件内容，第二个参数控制可见性
        )

        # 修改导出答案按钮绑定 - 增加文件下载组件作为输出
        export_answers_button.click(
            export_answers,
            inputs=[results, export_format],
            outputs=[status, download_file, download_file]  # 第一个参数控制文件内容，第二个参数控制可见性
        )

        # 删除按钮使用直接删除函数
        delete_button.click(
            direct_delete,
            inputs=[selected_question_id],  # 确保这里使用的是 selected_question_id
            outputs=[status, query_status, results, question_detail_group]
        )

        refresh_button.click(
            load_all_data,
            outputs=[query_status, results]
        )

        # 初始加载所有数据
        try:
            # 执行数据一致性检查
            db.verify_data_consistency()
            
            # 清理无效数据
            invalid_ids = []
            questions = db.get_all_error_questions(limit=100)
            for q in questions:
                if not db.get_error_question(q["id"]):
                    invalid_ids.append(q["id"])
                    
            if invalid_ids:
                logger.warning(f"发现 {len(invalid_ids)} 条无效错题记录，将进行清理")
                for invalid_id in invalid_ids:
                    db.delete_error_question(invalid_id)
            
            # 加载初始数据
            initial_status, initial_data = load_all_data()
            query_status.value = initial_status
            results.value = initial_data
            
            logger.info("错题查询页面初始化完成")
        except Exception as e:
            logger.error(f"加载初始数据失败: {str(e)}")
            logger.error(traceback.format_exc())
            query_status.value = f"加载初始数据失败: {str(e)}"

    return search_page
